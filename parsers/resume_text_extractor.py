import re
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from utils.logger import get_logger

logger = get_logger()

NOISE_REPLACEMENTS = {
    "\u2022": "-",
    "\u25cf": "-",
    "\u25aa": "-",
    "\u25a0": "-",
    "\u25cb": "-",
    "\u2013": "-",
    "\u2014": "-",
    "\u00a0": " ",
    "\u00e2\u20ac\u00a2": "-",
    "\u00e2\u20ac\u201c": "-",
    "\u00e2\u20ac\u201d": "-",
    "\u00e2\u2014\u2039": "-",
    "\u00e2\u201e\u00a2": "",
    "♂¶obile-alt": "Phone: ",
    "/envel⌢pe": "Email: ",
    "obile-alt": "Phone: ",
    "envel⌢pe": "Email: ",

}


TAB_PATTERN = re.compile(r"\t+")
SPACE_PATTERN = re.compile(r" +")
BLANK_LINE_PATTERN = re.compile(r"\n{3,}")
LINE_SPACE_PATTERN = re.compile(r"[^\S\n]+")
PAGE_NUMBER_PATTERN = re.compile(r"(?m)^\s*\d+\s*/\s*\d+\s*$")


SECTION_HEADINGS = [
    "summary",
    "profile",
    "profile summary",
    "skills",
    "technical skills",
    "experience",
    "work experience",
    "internships",
    "education",
    "projects",
    "certifications",
    "achievements",
    "contact",
    "personal details",
    "references",
]



def extract_text_from_pdf(file_path: str) -> str:
    text = ""

    try:
        reader = PdfReader(file_path)

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        logger.info(f"PDF text extracted successfully: {file_path}")
        return text

    except Exception as error:
        logger.error(f"PDF extraction failed: {file_path} | {error}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    text = ""

    try:
        document = Document(file_path)

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"

        logger.info(f"DOCX text extracted successfully: {file_path}")
        return text

    except Exception as error:
        logger.error(f"DOCX extraction failed: {file_path} | {error}")
        return ""


def clean_resume_text(raw_text: str) -> str:
    text = raw_text.replace("\r", "\n")

    for noisy_value, clean_value in NOISE_REPLACEMENTS.items():
        text = text.replace(noisy_value, clean_value)

    text = TAB_PATTERN.sub(" ", text)
    text = SPACE_PATTERN.sub(" ", text)
    text = BLANK_LINE_PATTERN.sub("\n\n", text)
    text = LINE_SPACE_PATTERN.sub(" ", text)
    text = PAGE_NUMBER_PATTERN.sub("", text)

    for heading in SECTION_HEADINGS:
        pattern = rf"(?im)^\s*{re.escape(heading)}\s*:?\s*$"
        text = re.sub(pattern, heading.upper(), text)

    return text.strip()


def extract_resume_text(file_path: str) -> str:
    file_extension = Path(file_path).suffix.lower()

    if file_extension == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        raw_text = extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported resume format: {file_extension}")
        return ""

    return clean_resume_text(raw_text)


def save_extracted_text(input_file: str, output_folder: str = "data/extracted_resumes") -> str:
    Path(output_folder).mkdir(parents=True, exist_ok=True)

    cleaned_text = extract_resume_text(input_file)

    output_file_name = Path(input_file).stem + ".txt"
    output_path = Path(output_folder) / output_file_name

    with open(output_path, "w", encoding="utf-8") as file:
        file.write(cleaned_text)

    logger.info(f"Cleaned resume text saved: {output_path}")

    return str(output_path)
